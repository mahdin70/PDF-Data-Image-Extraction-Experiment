import fitz 
from PIL import Image
from docx import Document
from docx.shared import Inches
import io
import os

def extract_text_and_images(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = []
    images = []

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        text.append(page.get_text("text"))

        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            images.append({
                "page": page_num,
                "image_index": img_index,
                "image": image
            })

    return text, images

def add_text_and_images_to_doc(text, images, doc_path):
    doc = Document()

    for page_text in text:
        doc.add_paragraph(page_text)
        doc.add_paragraph()

    for img in images:
        img_path = f"temp_image_{img['image_index']}.png"
        img["image"].save(img_path)
        doc.add_picture(img_path, width=Inches(4))
        os.remove(img_path)

    doc.save(doc_path)

def main():
    pdf_dir = 'pdfs'
    output_dir = 'output'
    pdf_name = 'PDF-2.pdf'
    pdf_path = os.path.join(pdf_dir, pdf_name)
    doc_name = pdf_name.replace('.pdf', '.docx')
    doc_path = os.path.join(output_dir, doc_name)

    text, images = extract_text_and_images(pdf_path)
    add_text_and_images_to_doc(text, images, doc_path)
    print(f"Document saved to: {doc_path}")


main()
